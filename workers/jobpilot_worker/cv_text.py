"""Plain-text extraction from CV files (pdf / docx / txt / md)."""
import io
from pathlib import Path


def cv_to_text(path: str | Path) -> str:
    p = Path(path)
    return cv_bytes_to_text(p.read_bytes(), p.name)


def cv_bytes_to_text(data: bytes, filename: str) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return _pdf(data)
    if suffix == ".docx":
        return _docx(data)
    if suffix in (".txt", ".md"):
        return data.decode("utf-8", errors="replace")
    raise ValueError(f"unsupported CV format '{suffix}' (use pdf, docx, txt, or md)")


def _pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _docx(data: bytes) -> str:
    import docx

    doc = docx.Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(p for p in parts if p and p.strip())
