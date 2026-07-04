"""Phase 2 plumbing: LiteLLM client contract and CV text extraction."""
import io

import httpx
import pytest

from jobpilot_worker import llm
from jobpilot_worker.cv_text import cv_bytes_to_text, cv_to_text


# ---- llm.embed ----

def _fake_embeddings_response(vectors):
    return {"data": [{"index": i, "embedding": v} for i, v in enumerate(vectors)]}


def test_embed_returns_vectors_in_input_order(monkeypatch):
    # response deliberately out of order — must be re-sorted by index
    def fake_post(url, headers=None, json=None, timeout=None):
        payload = {"data": [
            {"index": 1, "embedding": [1.0] * 1024},
            {"index": 0, "embedding": [0.0] * 1024},
        ]}
        return httpx.Response(200, json=payload, request=httpx.Request("POST", url))

    monkeypatch.setattr(llm.httpx, "post", fake_post)
    vecs = llm.embed(["a", "b"])
    assert vecs[0][0] == 0.0 and vecs[1][0] == 1.0


def test_embed_rejects_wrong_dimension(monkeypatch):
    def fake_post(url, headers=None, json=None, timeout=None):
        return httpx.Response(
            200, json=_fake_embeddings_response([[0.1] * 768]),
            request=httpx.Request("POST", url),
        )

    monkeypatch.setattr(llm.httpx, "post", fake_post)
    with pytest.raises(ValueError, match="768"):
        llm.embed(["a"])


def test_embed_empty_input_short_circuits():
    assert llm.embed([]) == []


def test_chat_returns_message_content(monkeypatch):
    def fake_post(url, headers=None, json=None, timeout=None):
        assert json["model"] == "cheap"
        return httpx.Response(
            200, json={"choices": [{"message": {"content": "hello"}}]},
            request=httpx.Request("POST", url),
        )

    monkeypatch.setattr(llm.httpx, "post", fake_post)
    assert llm.chat("hi") == "hello"


# ---- cv text extraction ----

def test_txt_cv_read_verbatim(tmp_path):
    f = tmp_path / "cv.txt"
    f.write_text("Ajinkya\nPython, SQL", encoding="utf-8")
    assert "Python, SQL" in cv_to_text(f)


def test_docx_cv_extracts_paragraphs_and_tables(tmp_path):
    import docx

    d = docx.Document()
    d.add_paragraph("Senior Engineer")
    table = d.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Skills"
    table.rows[0].cells[1].text = "Python"
    f = tmp_path / "cv.docx"
    d.save(str(f))

    text = cv_to_text(f)
    assert "Senior Engineer" in text
    assert "Python" in text


def test_bytes_variant_dispatches_on_filename(tmp_path):
    text = cv_bytes_to_text(b"plain text cv", "resume.txt")
    assert text == "plain text cv"


def test_unsupported_format_raises():
    with pytest.raises(ValueError, match="unsupported"):
        cv_bytes_to_text(b"", "resume.pages")
